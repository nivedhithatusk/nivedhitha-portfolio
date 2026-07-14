"""Update resume DOCX: title, summary, and technical skills."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Nivedhitha_Professional_Resume_Updated_v3.docx"
PUBLIC_DOCX = ROOT / "public" / "Nivedhitha_Professional_Resume.docx"

SKILL_ROWS = [
    (
        "Deployment & Infrastructure: ",
        "Application Deployment, Linux Server Administration, Production Support, Server Migration, Infrastructure Management, Domain & DNS Management, SSL Certificate Installation, Reverse Proxy Configuration",
    ),
    (
        "Cloud & DevOps: ",
        "AWS EC2, Amazon S3, Amazon CloudFront, Amazon Route 53, AWS Certificate Manager (ACM), Elastic Load Balancer (ELB), AWS WAF, Nginx, PM2, Ubuntu, cPanel, GoDaddy Hosting",
    ),
    (
        "Frontend: ",
        "React.js, Next.js, HTML5, CSS3, JavaScript (ES6+), Tailwind CSS, JSP, AJAX",
    ),
    (
        "Backend: ",
        "Node.js, Java, Core Java, J2EE, Spring Boot, Hibernate, Struts, REST APIs, JSON, JDBC",
    ),
    (
        "Databases: ",
        "MySQL, PostgreSQL, Oracle Database",
    ),
    (
        "Version Control: ",
        "Git, GitHub",
    ),
    (
        "Tools & IDEs: ",
        "VS Code, Cursor, IntelliJ IDEA, Spring Tool Suite (STS), MyEclipse, Postman, DBeaver, Adminer, SonarQube, OpenProject, MobaXterm, Electerm",
    ),
    (
        "Project Management & Business Tools: ",
        "Asana, Workforce, Workcloud, SOS",
    ),
]

SUMMARY = (
    "Software Engineer with 6+ years of experience in designing, developing, and "
    "supporting enterprise web applications, including 2+ years of hands-on experience "
    "in application deployment, cloud infrastructure, and production support. "
    "Proficient in Java, Spring Boot, Node.js, React.js, and Next.js, with practical "
    "experience deploying and maintaining applications on Linux and AWS environments. "
    "Skilled in configuring Nginx, PM2, Amazon EC2, Route 53, SSL, Load Balancers, WAF, "
    "cPanel, DNS management, and production troubleshooting. Experienced in delivering "
    "secure, scalable, and reliable solutions while collaborating with cross-functional "
    "teams and clients across banking, fintech, healthcare, media, and enterprise domains."
)


def set_run_text(paragraph: Paragraph, text: str, *, bold: bool | None = None) -> None:
    if not paragraph.runs:
        run = paragraph.add_run(text)
        if bold is not None:
            run.bold = bold
        return
    paragraph.runs[0].text = text
    if bold is not None:
        paragraph.runs[0].bold = bold
    for run in paragraph.runs[1:]:
        run.text = ""


def set_labeled_skill(paragraph: Paragraph, label: str, value: str) -> None:
    # Keep/reuse run formatting: bold label + normal value
    if len(paragraph.runs) >= 2:
        paragraph.runs[0].text = label
        paragraph.runs[0].bold = True
        paragraph.runs[1].text = value
        paragraph.runs[1].bold = False
        for run in paragraph.runs[2:]:
            run.text = ""
        return
    if paragraph.runs:
        paragraph.runs[0].text = ""
    label_run = paragraph.add_run(label)
    label_run.bold = True
    value_run = paragraph.add_run(value)
    value_run.bold = False


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def clear_runs(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def main() -> None:
    doc = Document(str(DOCX_PATH))
    paras = doc.paragraphs

    # Title line
    set_run_text(paras[1], "Software Engineer", bold=True)

    # Summary heading already correct; ensure body text
    set_run_text(paras[6], SUMMARY)

    # Skills heading
    set_run_text(paras[7], "TECHNICAL SKILLS", bold=True)

    existing_skill_paras = paras[8:15]  # currently 7 skill lines
    needed = len(SKILL_ROWS)

    # Ensure we have enough paragraphs for all skill rows
    skill_paras = list(existing_skill_paras)
    while len(skill_paras) < needed:
        skill_paras.append(insert_paragraph_after(skill_paras[-1]))

    # If somehow more exist (future), clear extras later by blanking
    for i, (label, value) in enumerate(SKILL_ROWS):
        set_labeled_skill(skill_paras[i], label, value)

    for extra in skill_paras[needed:]:
        clear_runs(extra)

    doc.save(str(DOCX_PATH))
    PUBLIC_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(PUBLIC_DOCX))
    print(f"Updated: {DOCX_PATH}")
    print(f"Copied:  {PUBLIC_DOCX}")


if __name__ == "__main__":
    main()
