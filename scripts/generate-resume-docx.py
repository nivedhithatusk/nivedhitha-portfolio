"""Generate a polished, ATS-friendly professional resume DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
DOCX_ROOT = ROOT / "Nivedhitha_Professional_Resume_Updated_v3.docx"
DOCX_PUBLIC = ROOT / "public" / "Nivedhitha_Professional_Resume.docx"

# Brand colors (match existing resume style)
INK = RGBColor(0x0F, 0x17, 0x2A)
TEAL = RGBColor(0x0D, 0x94, 0x88)
MUTED = RGBColor(0x47, 0x55, 0x69)
RULE = "CBD5E1"

LINKEDIN = "https://www.linkedin.com/in/nivedhitha-praba-07925913a"
PORTFOLIO = "https://nivedhitha-portfolio.vercel.app"
EMAIL = "nivedhithapraba@gmail.com"

SUMMARY = (
    "Full Stack Developer & DevOps Engineer with 6+ years of experience designing, "
    "developing, and supporting enterprise web applications, including 2+ years of "
    "hands-on application deployment, cloud infrastructure, and production support. "
    "Proficient in Java, Spring Boot, Node.js, React.js, and Next.js, with practical "
    "experience deploying and maintaining applications on Linux and AWS. Skilled in "
    "Nginx, PM2, Amazon EC2, Route 53, SSL, Load Balancers, WAF, cPanel, DNS management, "
    "and production troubleshooting. Delivers secure, scalable solutions in collaboration "
    "with cross-functional teams and clients across banking, fintech, healthcare, media, "
    "and enterprise domains."
)

SKILL_ROWS = [
    (
        "Deployment",
        "Application Deployment, Linux Admin, Production Support, Server Migration, "
        "Infrastructure Management, Domain & DNS, SSL, Reverse Proxy",
    ),
    (
        "Cloud & DevOps",
        "AWS EC2, Lightsail, S3, CloudFront, Route 53, ACM, ELB, WAF, Nginx, PM2, "
        "Ubuntu, cPanel, CWP, GoDaddy Hosting",
    ),
    (
        "Frontend",
        "React.js, Next.js, HTML5, CSS3, JavaScript (ES6+), Tailwind CSS, JSP, AJAX",
    ),
    (
        "Backend",
        "Node.js, Java, J2EE, Spring Boot, Hibernate, Struts, PHP, Strapi, REST APIs, "
        "JSON, JDBC",
    ),
    (
        "Databases",
        "MySQL, PostgreSQL, MongoDB, Oracle Database",
    ),
    (
        "Tools",
        "Git, GitHub, VS Code, Cursor, IntelliJ IDEA, STS, MyEclipse, Postman, "
        "DBeaver, Adminer, SonarQube, OpenProject, MobaXterm, Electerm, Asana, "
        "Workforce, workspace",
    ),
]


def set_run_font(run, *, size=10.5, bold=False, color=INK, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_horizontal_rule(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text: str, url: str, *, size=10.5, color=MUTED) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    rPr.append(color_el)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_paragraph(
    paragraph,
    *,
    space_before=0,
    space_after=4,
    line_spacing=1.08,
    align=None,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, space_before=12, space_after=4, line_spacing=1.05)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=TEAL)
    add_horizontal_rule(p)


def add_body(doc: Document, text: str, *, space_after=6) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, space_after=space_after, line_spacing=1.12)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_skill_row(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, space_after=2, line_spacing=1.08)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True, color=INK)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10.5, color=INK)


def set_cell_border(cell, **kwargs) -> None:
    """kwargs: top/left/bottom/right -> {val, sz, color}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), edge_data.get("val", "single"))
        element.set(qn("w:sz"), str(edge_data.get("sz", 4)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), edge_data.get("color", RULE))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_skills_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Compact 2-column skills grid — neat alignment, fewer wasted lines."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False

    # ~1.55" category | ~5.4" skills (A4 content width ~7")
    widths = (Inches(1.55), Inches(5.4))
    hairline = {"val": "single", "sz": 4, "color": RULE}

    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        cat_cell, skill_cell = row.cells

        cat_cell.width = widths[0]
        skill_cell.width = widths[1]

        # Category cell
        cat_cell.text = ""
        cp = cat_cell.paragraphs[0]
        style_paragraph(cp, space_before=1, space_after=1, line_spacing=1.05)
        cr = cp.add_run(label)
        set_run_font(cr, size=9.5, bold=True, color=TEAL)
        set_cell_shading(cat_cell, "F0FDFA")  # mild teal tint
        set_cell_border(cat_cell, top=hairline, left=hairline, bottom=hairline, right=hairline)

        # Skills cell
        skill_cell.text = ""
        sp = skill_cell.paragraphs[0]
        style_paragraph(sp, space_before=1, space_after=1, line_spacing=1.05)
        sr = sp.add_run(value)
        set_run_font(sr, size=9.5, color=INK)
        set_cell_border(skill_cell, top=hairline, left=hairline, bottom=hairline, right=hairline)

        # Compact cell margins
        for cell in (cat_cell, skill_cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for m in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{m}")
                node.set(qn("w:w"), "40" if m in ("left", "right") else "28")
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)

    # Tight space after table
    spacer = doc.add_paragraph()
    style_paragraph(spacer, space_before=0, space_after=2, line_spacing=1.0)



def add_job_header(doc: Document, company_role: str, period_location: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, space_before=8, space_after=0, line_spacing=1.05)
    run = p.add_run(company_role)
    set_run_font(run, size=10.5, bold=True, color=INK)

    p2 = doc.add_paragraph()
    style_paragraph(p2, space_before=0, space_after=2, line_spacing=1.05)
    run2 = p2.add_run(period_location)
    set_run_font(run2, size=10, color=MUTED)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, space_after=1, line_spacing=1.08)
    # Clear default run if any, then add styled text
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10.5, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=INK)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, space_before=3, space_after=1, line_spacing=1.05)
    run = p.add_run(text)
    set_run_font(run, size=10, bold=True, color=MUTED)


def build() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4-ish / letter-compatible width used before
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # Name
    name = doc.add_paragraph()
    style_paragraph(name, space_after=2, line_spacing=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = name.add_run("NIVEDHITHA A")
    set_run_font(r, size=22, bold=True, color=INK)

    # Title
    title = doc.add_paragraph()
    style_paragraph(title, space_after=2, line_spacing=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = title.add_run("Full Stack Developer & DevOps Engineer")
    set_run_font(r, size=12, bold=True, color=TEAL)

    # Contact
    contact = doc.add_paragraph()
    style_paragraph(contact, space_after=1, line_spacing=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    c1 = contact.add_run(f"Chennai, Tamil Nadu  ·  {EMAIL}")
    set_run_font(c1, size=10.5, color=MUTED)

    links = doc.add_paragraph()
    style_paragraph(links, space_after=8, line_spacing=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_hyperlink(links, "LinkedIn", LINKEDIN)
    sep = links.add_run("  ·  ")
    set_run_font(sep, size=10.5, color=MUTED)
    add_hyperlink(links, "Portfolio", PORTFOLIO)

    # Summary
    add_heading(doc, "Professional Summary")
    add_body(doc, SUMMARY, space_after=6)

    # Skills — compact organized table (not long wrapping paragraphs)
    add_heading(doc, "Technical Skills")
    add_skills_table(doc, SKILL_ROWS)

    # Experience
    add_heading(doc, "Professional Experience")

    # Tuskmelon
    add_job_header(
        doc,
        "Tuskmelon Business Solutions  |  Senior Full Stack Developer & DevOps Engineer",
        "Sep 2023 – Present  |  Chennai",
    )
    for bullet in [
        "Develop and deploy full-stack applications using React.js, Next.js, Node.js, PHP, and Strapi.",
        "Provision and maintain AWS EC2, Amazon Lightsail, GoDaddy, cPanel, and CWP environments, including SSL and domain configuration.",
        "Manage Nginx + PM2 deployments, AWS Certificate Manager (ACM), and DNS via Route 53 and GoDaddy.",
        "Handle multi-environment server configuration, logging, and production troubleshooting for banking, healthcare, and enterprise clients.",
    ]:
        add_bullet(doc, bullet)

    add_subheading(doc, "Key Projects")
    for bullet in [
        "RHFL (Repco Home) — Next.js frontend and Strapi + MySQL backend deployed on Linux.",
        "Equitas Gurukul — Next.js frontend/admin with Strapi + MySQL; deployed on EC2 (Ubuntu).",
        "Medall Healthcare — Next.js frontend and Node.js + MongoDB backend on EC2 (Ubuntu).",
        "City Union Bank (GMB) — React.js + Node.js + MySQL; EC2 deployment with GoDaddy domain.",
        "Equitas Locate — Next.js + Strapi + MySQL on AWS with WAF, Load Balancer, ACM, Security Groups, Nginx, and PM2.",
        "Repco Bank & Repco Bank (GMB) — PHP/cPanel website and React/Node.js GMB stack on Ubuntu/EC2.",
        "Internal platforms (Workforce, Workspace, TuskQR, Tusk Cloud, Social Media Manager) — Next.js/React/Node.js + MySQL on EC2 and Lightsail.",
        "Uniscan — Admin panel and website forms developed and deployed on Ubuntu; ongoing production support for Unico, RHFL, Equitas, CUB, and Repco applications.",
    ]:
        add_bullet(doc, bullet)

    # CreditMantri
    add_job_header(
        doc,
        "CreditMantri  |  Full Stack Developer",
        "Sep 2022 – Mar 2023  |  Chennai",
    )
    for bullet in [
        "Led HDFC Fintech PL & CC API development for loan processing, customer onboarding, and KYC document uploads.",
        "Implemented webhook integrations for SMS, Email, and WhatsApp via WebEngage.",
        "Built custom SMS APIs for the Kurundhagaval project with bulk and single-read processing logic.",
        "Delivered mission-critical SBI modules (LDB, EDB, Emudhra) in coordination with QA and DevOps teams.",
        "Owned end-to-end API delivery, Postman testing, database schema design, and production deployments.",
    ]:
        add_bullet(doc, bullet)

    # HTC
    add_job_header(
        doc,
        "HTC Global Services  |  L2 Engineer",
        "Feb 2022 – Sep 2022  |  Chennai",
    )
    for bullet in [
        "Completed a structured Java Full Stack Development program covering Spring Boot, REST APIs, and relational databases.",
        "Built applications using Java, Spring Boot, JSP, MySQL, PostgreSQL, and Oracle Database.",
        "Gained practical exposure to Git, Linux fundamentals, and cross-platform deployment practices.",
    ]:
        add_bullet(doc, bullet)

    # Soft Media
    add_job_header(
        doc,
        "Soft Media ERP  |  Full Stack Developer",
        "Jun 2019 – Feb 2022  |  Chennai",
    )
    for bullet in [
        "Developed the Daily Thanthi Media ERP covering Advertisement, Circulation, Accounts, HR, and Purchase modules.",
        "Implemented DAO design patterns to maintain data integrity across distributed systems.",
        "Built dynamic PDF/Text reporting modules for business stakeholders.",
        "Automated payroll and procurement workflows with role-based access controls.",
    ]:
        add_bullet(doc, bullet)

    # Education
    add_heading(doc, "Education")
    add_body(
        doc,
        "Master of Computer Applications (MCA) — University of Madras  |  2018 – 2021",
        space_after=2,
    )
    add_body(
        doc,
        "Bachelor of Computer Applications (BCA) — St. Anne's Arts & Science College  |  2015 – 2018",
        space_after=0,
    )

    return doc


def main() -> None:
    doc = build()
    DOCX_PUBLIC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_ROOT))
    doc.save(str(DOCX_PUBLIC))
    print(f"Wrote: {DOCX_ROOT}")
    print(f"Wrote: {DOCX_PUBLIC}")


if __name__ == "__main__":
    main()
